from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Student Attendance Management System\nAPI Endpoints Documentation')
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph('Generated for the current backend project structure.')
doc.add_paragraph('')

sections = [
    ('Base / General', [
        ('GET /', 'Root endpoint returning a simple API status message.'),
        ('GET /admin/', 'Django admin panel.'),
        ('GET /api-auth/', 'Django REST Framework browsable API auth interface.'),
    ]),
    ('Accounts', [
        ('POST /api/accounts/login/', 'Login endpoint for users.'),
        ('GET /api/accounts/me/', 'Get current authenticated user details.'),
        ('GET /api/accounts/users/', 'List users.'),
        ('POST /api/accounts/users/', 'Create user.'),
        ('GET /api/accounts/users/{id}/', 'Retrieve user.'),
        ('PUT /api/accounts/users/{id}/', 'Update user.'),
        ('DELETE /api/accounts/users/{id}/', 'Delete user.'),
        ('GET /api/accounts/teachers/', 'List teachers.'),
        ('POST /api/accounts/teachers/', 'Create teacher.'),
        ('GET /api/accounts/teachers/{id}/', 'Retrieve teacher.'),
        ('PUT /api/accounts/teachers/{id}/', 'Update teacher.'),
        ('DELETE /api/accounts/teachers/{id}/', 'Delete teacher.'),
        ('GET /api/accounts/students/', 'List students.'),
        ('POST /api/accounts/students/', 'Create student.'),
        ('GET /api/accounts/students/{id}/', 'Retrieve student.'),
        ('PUT /api/accounts/students/{id}/', 'Update student.'),
        ('DELETE /api/accounts/students/{id}/', 'Delete student.'),
    ]),
    ('Courses', [
        ('GET /api/courses/courses/', 'List courses.'),
        ('POST /api/courses/courses/', 'Create course.'),
        ('GET /api/courses/courses/{id}/', 'Retrieve course.'),
        ('PUT /api/courses/courses/{id}/', 'Update course.'),
        ('DELETE /api/courses/courses/{id}/', 'Delete course.'),
    ]),
    ('Batches', [
        ('GET /api/batches/batches/', 'List batches.'),
        ('POST /api/batches/batches/', 'Create batch.'),
        ('GET /api/batches/batches/{id}/', 'Retrieve batch.'),
        ('PUT /api/batches/batches/{id}/', 'Update batch.'),
        ('DELETE /api/batches/batches/{id}/', 'Delete batch.'),
    ]),
    ('Schedules', [
        ('GET /api/schedules/my-schedules/', 'Get schedules for the logged-in teacher.'),
        ('GET /api/schedules/', 'List schedules.'),
        ('POST /api/schedules/', 'Create schedule.'),
        ('GET /api/schedules/{id}/', 'Retrieve schedule.'),
        ('PUT /api/schedules/{id}/', 'Update schedule.'),
        ('DELETE /api/schedules/{id}/', 'Delete schedule.'),
    ]),
    ('Attendance', [
        ('GET /api/attendance/sessions/', 'List attendance sessions.'),
        ('POST /api/attendance/sessions/', 'Create attendance session.'),
        ('GET /api/attendance/sessions/{id}/', 'Retrieve attendance session.'),
        ('PUT /api/attendance/sessions/{id}/', 'Update attendance session.'),
        ('DELETE /api/attendance/sessions/{id}/', 'Delete attendance session.'),
        ('GET /api/attendance/records/', 'List attendance records.'),
        ('POST /api/attendance/records/', 'Create attendance record.'),
        ('GET /api/attendance/records/{id}/', 'Retrieve attendance record.'),
        ('PUT /api/attendance/records/{id}/', 'Update attendance record.'),
        ('DELETE /api/attendance/records/{id}/', 'Delete attendance record.'),
        ('POST /api/attendance/mark/', 'Mark attendance for one or more students.'),
        ('GET /api/attendance/schedule/{schedule_id}/students/', 'Get students for a given schedule.'),
        ('GET /api/attendance/session/{session_id}/', 'Get attendance history for a session.'),
    ]),
    ('Dashboard', [
        ('GET /api/dashboard/admin/', 'Admin dashboard metrics.'),
        ('GET /api/dashboard/teacher/', 'Teacher dashboard metrics.'),
        ('GET /api/dashboard/student/', 'Student dashboard metrics.'),
    ]),
]

for title, items in sections:
    doc.add_heading(title, level=1)
    for method, desc in items:
        p = doc.add_paragraph()
        p.add_run(method).bold = True
        p.add_run(' - ' + desc)

out_path = 'API_Endpoints_Documentation.docx'
doc.save(out_path)
print(out_path)
